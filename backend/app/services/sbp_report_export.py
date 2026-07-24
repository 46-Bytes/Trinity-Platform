"""
Strategic Business Plan Report Exporter
Generates professional .docx files from the assembled plan.
"""
import io
import logging
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup

from app.models.strategic_business_plan import StrategicBusinessPlan

logger = logging.getLogger(__name__)

SECTION_INTROS = {
    "executive_summary": (
        "This executive summary provides a high-level overview of the strategic direction, "
        "key priorities, and anticipated outcomes for the planning period."
    ),
    "strategic_intent": (
        "This section defines the long-term intent of the business. It establishes the "
        "directional anchors that guide all strategic, operational, and investment decisions "
        "over the planning horizon."
    ),
    "business_context": (
        "This section provides the operating context for the business — its history, model, "
        "market position, and the key forces shaping its environment."
    ),
    "external_internal_analysis": (
        "This section analyses the external environment and internal capability of the business "
        "to surface strategic tensions, opportunities, and risks."
    ),
    "key_resources_capabilities": (
        "This section identifies the key resources and capabilities that underpin the business's "
        "competitive position and strategic options."
    ),
    "customer_dynamics": (
        "This section examines customer segments, buying behaviour, and the relationships that "
        "drive revenue and loyalty."
    ),
    "growth_opportunities": (
        "This section identifies and evaluates the primary growth opportunities available to "
        "the business over the planning horizon."
    ),
    "operations_strategy": (
        "This section defines the operational priorities and key recommendations required to "
        "deliver the strategy efficiently and at scale."
    ),
    "hr_strategy": (
        "This section outlines the people and leadership strategy required to build the "
        "capability, culture, and capacity needed to execute the plan."
    ),
    "marketing_sales_strategy": (
        "This section defines the marketing, sales, and brand approach required to attract, "
        "convert, and retain the target customer."
    ),
    "financial_overview": (
        "This section provides the financial framework for the strategy — including projected "
        "performance, investment requirements, and capital discipline."
    ),
    "risk_matrix": (
        "This section identifies and evaluates the key risks to strategy execution, together "
        "with mitigation actions and risk owners."
    ),
    "actions_next_steps": (
        "This section sets out the prioritised action plan and implementation roadmap that "
        "will translate strategy into execution."
    ),
    "strategic_alignment": (
        "This section integrates the key strategic implications across all chapters and sets "
        "out the overarching commitment to execution and the way forward."
    ),
}


class SBPReportExporter:
    """Generates a formatted .docx Strategic Business Plan document."""

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        heading = doc.add_heading(text, level=level)
        color = RGBColor(0x1A, 0x1A, 0x1A) if level == 1 else RGBColor(0x1F, 0x4D, 0x78)
        for run in heading.runs:
            run.font.color.rgb = color
            run.font.name = "Calibri"

    def _add_toc(self, doc: Document):
        """Insert a real Word TOC field — auto-updates on open in Word."""
        self._add_heading(doc, "Table of Contents", level=1)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\h \\o "1-2" \\u '
        fldChar_sep = OxmlElement('w:fldChar')
        fldChar_sep.set(qn('w:fldCharType'), 'separate')
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_sep)
        run._r.append(fldChar_end)

    def _add_header(self, doc: Document, client_name: str):
        """Add a header with the client name and a navy bottom border."""
        section = doc.sections[0]
        header = section.header
        para = header.paragraphs[0]
        run = para.add_run(f"{client_name} — Strategic Business Plan")
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F4D78')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_footer(self, doc: Document):
        """Add a centered footer with advisory credit and live page number."""
        section = doc.sections[0]
        footer = section.footer
        para = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = para.add_run("Prepared by Benchmark Business Advisory  |  Page ")
        run1.font.size = Pt(8)
        run1.font.name = "Calibri"
        run1.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run2 = para.add_run()
        run2._r.append(fldChar1)
        run2._r.append(instrText)
        run2._r.append(fldChar2)
        run2.font.size = Pt(8)
        run2.font.name = "Calibri"
        run2.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    def _add_section_intro(self, doc: Document, key: str):
        """Insert an italic gray framing paragraph after each chapter heading."""
        text = SECTION_INTROS.get(key)
        if not text:
            return
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        p.paragraph_format.space_after = Pt(10)

    def _apply_body_spacing(self, para):
        """Set 1.5× line spacing and 7pt after-paragraph spacing on a body paragraph."""
        pPr = para._p.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        spacing.set(qn('w:after'), '140')   # 7pt
        spacing.set(qn('w:line'), '360')    # 1.5× (single = 240)
        spacing.set(qn('w:lineRule'), 'auto')

    def _set_cell_shading(self, cell, fill_hex: str):
        """Set a table cell's background colour via OOXML shading."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        tcPr.append(shd)

    def _html_to_docx(self, doc: Document, html_content: str, skip_first_heading: bool = False):
        """Convert HTML content to docx paragraphs."""
        if not html_content:
            return

        soup = BeautifulSoup(html_content, "html.parser")
        first_heading_skipped = not skip_first_heading

        for element in soup.children:
            if element.name in ("h1", "h2", "h3", "h4"):
                if not first_heading_skipped:
                    first_heading_skipped = True
                    continue  # drop redundant repeat of the chapter title
                level_map = {"h1": 1, "h2": 2, "h3": 3, "h4": 4}
                self._add_heading(doc, element.get_text(strip=True), level=level_map.get(element.name, 3))
            elif element.name == "p":
                para = doc.add_paragraph(element.get_text(strip=True))
                para.style.font.size = Pt(11)
                self._apply_body_spacing(para)
            elif element.name in ("ul", "ol"):
                for li in element.find_all("li", recursive=False):
                    para = doc.add_paragraph(li.get_text(strip=True), style="List Bullet")
                    self._apply_body_spacing(para)
            elif element.name == "table":
                self._html_table_to_docx(doc, element)
            elif element.name == "blockquote":
                para = doc.add_paragraph(element.get_text(strip=True))
                para.paragraph_format.left_indent = Cm(1.5)
                para.style.font.italic = True
                self._apply_body_spacing(para)
            elif hasattr(element, "get_text"):
                text = element.get_text(strip=True)
                if text:
                    para = doc.add_paragraph(text)
                    self._apply_body_spacing(para)

    def _html_table_to_docx(self, doc: Document, table_element):
        """Convert an HTML table to a docx table."""
        rows_data = []
        for tr in table_element.find_all("tr"):
            cells = [td.get_text(strip=True) for td in tr.find_all(["th", "td"])]
            if cells:
                rows_data.append(cells)

        if not rows_data:
            return

        max_cols = max(len(row) for row in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=max_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row_data in enumerate(rows_data):
            for j, cell_text in enumerate(row_data):
                if j < max_cols:
                    cell = table.cell(i, j)
                    cell.text = cell_text
                    if i == 0:
                        self._set_cell_shading(cell, '1A1A1A')
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.font.name = "Calibri"

        doc.add_paragraph("")

    def generate_docx(self, plan: StrategicBusinessPlan) -> bytes:
        """Generate the main Strategic Business Plan .docx document."""
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        year = datetime.now().year
        client_name = plan.client_name or "Client"

        # ── Cover page ──────────────────────────────────────────────────
        for i, word in enumerate(["STRATEGIC", "BUSINESS", "PLAN"]):
            p = doc.add_paragraph()
            run = p.add_run(word)
            run.bold = True
            run.font.size = Pt(36)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            p.paragraph_format.space_before = Pt(120) if i == 0 else Pt(10)
            p.paragraph_format.space_after = Pt(10)

        p = doc.add_paragraph()
        run = p.add_run(client_name.upper())
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0xC9, 0xA2, 0x27)
        p.paragraph_format.space_before = Pt(30)
        p.paragraph_format.space_after = Pt(5)

        horizon = plan.planning_horizon or "3-Year"
        p = doc.add_paragraph()
        run = p.add_run(f"{horizon} Planning Horizon | FY{year}–FY{year + 2}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        p.paragraph_format.space_after = Pt(5)

        p = doc.add_paragraph()
        run = p.add_run("Prepared by Benchmark Business Advisory")
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

        doc.add_page_break()

        # ── Header and footer ────────────────────────────────────────────
        self._add_header(doc, client_name)
        self._add_footer(doc)

        # ── Table of Contents ────────────────────────────────────────────
        self._add_toc(doc)
        doc.add_page_break()

        # ── Sections ─────────────────────────────────────────────────────
        final_plan = plan.final_plan or {}
        sections = final_plan.get("sections", [])
        numbered_index = 0

        for section in sections:
            if not section.get("content"):
                continue

            key = section.get("key", "")

            if key == "executive_summary":
                self._add_heading(doc, "Executive Summary", level=1)
            else:
                numbered_index += 1
                self._add_heading(doc, f"{numbered_index}. {section['title']}", level=1)

            self._add_section_intro(doc, key)
            self._html_to_docx(doc, section["content"], skip_first_heading=True)
            doc.add_page_break()

        # ── Save ──────────────────────────────────────────────────────────
        buffer = io.BytesIO()
        doc.save(buffer)
        doc_bytes = buffer.getvalue()
        logger.info(f"Generated SBP report for plan {plan.id} ({len(doc_bytes)} bytes)")

        plan.status = "completed"
        plan.completed_at = datetime.now()

        return doc_bytes


class SBPEmployeeExporter(SBPReportExporter):
    """Generates an employee-facing strategy document variant."""

    def generate_employee_docx(self, plan: StrategicBusinessPlan) -> bytes:
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        title = doc.add_heading("Our Strategic Direction", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(plan.client_name or "Our Company")
        run.font.size = Pt(16)

        doc.add_page_break()

        if plan.employee_plan:
            # Use advisor-edited employee plan
            sections = plan.employee_plan.get("sections", [])
            for section in sections:
                if section.get("included", True) and section.get("content"):
                    self._add_heading(doc, section["title"], level=1)
                    self._html_to_docx(doc, section["content"])
                    doc.add_page_break()
        else:
            # Fall back to filtering from final_plan
            final_plan = plan.final_plan or {}
            sections = final_plan.get("sections", [])
            employee_keys = [
                "executive_summary",
                "strategic_intent",
                "growth_opportunities",
                "operations_strategy",
                "hr_strategy",
                "marketing_sales_strategy",
            ]
            for section in sections:
                if section.get("key") in employee_keys and section.get("content"):
                    self._add_heading(doc, section["title"], level=1)
                    self._html_to_docx(doc, section["content"])
                    doc.add_page_break()

        buffer = io.BytesIO()
        doc.save(buffer)
        doc_bytes = buffer.getvalue()
        logger.info(f"Generated employee strategy document for plan {plan.id} ({len(doc_bytes)} bytes)")

        return doc_bytes


def get_sbp_report_exporter() -> SBPReportExporter:
    return SBPReportExporter()


def get_sbp_employee_exporter() -> SBPEmployeeExporter:
    return SBPEmployeeExporter()
