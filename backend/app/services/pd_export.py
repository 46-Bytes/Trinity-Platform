"""
Position Description Word exporter.

Builds the .docx from the approved PD content, following the seven-section
structure of the sample position descriptions.
"""
from __future__ import annotations

from typing import Any, Dict, List, Optional
import io
import logging

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

BODY_FONT = "Calibri"
BODY_SIZE = Pt(11)

# Sections 1-6 always render. Transition Focus is section 7 and only appears
# where the role carries handover items.
SECTION_TITLES = {
    "position_purpose": "Position Purpose",
    "key_responsibilities": "Key Responsibilities",
    "decision_making_authority": "Decision-Making Authority",
    "key_relationships": "Key Relationships",
    "kpis": "Key Performance Indicators (KPIs)",
    "behavioural_expectations": "Behavioural Expectations",
    "transition_focus": "Transition Focus",
}


class PositionDescriptionExporter:
    """Generates the Position Description document for a single role."""

    def generate_document_bytes(
        self,
        role_title: str,
        pd_content: Dict[str, Any],
        client_name: Optional[str] = None,
        fy_range: Optional[str] = None,
    ) -> io.BytesIO:
        """
        Build the .docx for one role.

        Args:
            role_title: The role the PD describes.
            pd_content: The approved PD sections.
            client_name: Business name shown under the title.
            fy_range: Financial year range appended to the Transition Focus heading.

        Returns:
            BytesIO stream positioned at the start, ready to stream to the client.
        """
        content = pd_content or {}

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = BODY_FONT
        style.font.size = BODY_SIZE

        self._add_title_block(doc, role_title, client_name)

        number = 1

        purpose = self._clean(content.get("position_purpose"))
        if purpose:
            self._add_section_heading(doc, number, SECTION_TITLES["position_purpose"])
            doc.add_paragraph(purpose)
            number += 1

        themes = content.get("key_responsibilities") or []
        if themes:
            self._add_section_heading(doc, number, SECTION_TITLES["key_responsibilities"])
            self._add_themed_responsibilities(doc, themes)
            number += 1

        for key in ("decision_making_authority", "key_relationships", "kpis", "behavioural_expectations"):
            items = self._clean_list(content.get(key))
            if not items:
                continue
            self._add_section_heading(doc, number, SECTION_TITLES[key])
            self._add_bullets(doc, items)
            number += 1

        transition = self._clean_list(content.get("transition_focus"))
        if transition:
            heading = SECTION_TITLES["transition_focus"]
            if fy_range:
                heading = f"{heading} {fy_range}"
            self._add_section_heading(doc, number, heading)
            self._add_bullets(doc, transition)

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        logger.info(f"Generated position description document for '{role_title}'")
        return stream

    # ==================== Building blocks ====================

    @staticmethod
    def _add_title_block(doc: Document, role_title: str, client_name: Optional[str]) -> None:
        """Title, role and business name, matching the sample PDs."""
        title = doc.add_heading("Position Description", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        role = doc.add_paragraph()
        role.alignment = WD_ALIGN_PARAGRAPH.CENTER
        role_run = role.add_run(role_title)
        role_run.bold = True
        role_run.font.size = Pt(14)

        if client_name:
            client = doc.add_paragraph()
            client.alignment = WD_ALIGN_PARAGRAPH.CENTER
            client_run = client.add_run(client_name)
            client_run.font.size = Pt(12)

    @staticmethod
    def _add_section_heading(doc: Document, number: int, text: str) -> None:
        """Numbered section heading, e.g. '1. Position Purpose'."""
        doc.add_heading(f"{number}. {text}", level=1)

    @staticmethod
    def _add_bullets(doc: Document, items: List[str]) -> None:
        """Write one bullet per item."""
        for item in items:
            doc.add_paragraph(item, style="List Bullet")

    def _add_themed_responsibilities(self, doc: Document, themes: List[Any]) -> None:
        """Write each theme as a bold sub-heading followed by its bullets."""
        for theme in themes or []:
            if not isinstance(theme, dict):
                continue

            name = self._clean(theme.get("theme"))
            responsibilities = self._clean_list(theme.get("responsibilities"))
            if not responsibilities:
                continue

            if name:
                heading = doc.add_paragraph()
                run = heading.add_run(name)
                run.bold = True

            self._add_bullets(doc, responsibilities)

    # ==================== Normalisation ====================

    @staticmethod
    def _clean(value: Any) -> Optional[str]:
        """Normalise a value to trimmed text, or None when blank."""
        if value is None:
            return None
        text = str(value).strip()
        return text or None

    @classmethod
    def _clean_list(cls, values: Any) -> List[str]:
        """Normalise a list, dropping blanks so empty sections are skipped."""
        if not isinstance(values, list):
            return []
        cleaned = [cls._clean(value) for value in values]
        return [value for value in cleaned if value]


_exporter: Optional[PositionDescriptionExporter] = None


def get_position_description_exporter() -> PositionDescriptionExporter:
    """Get the shared position description exporter instance"""
    global _exporter
    if _exporter is None:
        _exporter = PositionDescriptionExporter()
    return _exporter
