"""
BBA Report Export Service
Generates Word documents from BBA report data using python-docx
"""
from typing import Optional
from datetime import datetime
import io
import logging
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.models.bba import BBA

logger = logging.getLogger(__name__)


class BBAReportExporter:
    """
    Exports BBA diagnostic reports to Word document format.
    
    Creates professional reports following BBA's house style with:
    - Title page
    - Executive summary
    - Key findings & recommendations snapshot table
    - Key findings ranked by importance
    - 12-month recommendations plan
    - Timeline summary
    """
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for Word export. "
                "Install it with: pip install python-docx"
            )
    
    def generate_report(self, bba: BBA) -> bytes:
        """
        Generate a Word document from BBA data.
        
        Args:
            bba: BBA model with all report data
            
        Returns:
            Bytes of the generated Word document
        """
        logger.info(f"[BBA Export] Generating Word report for BBA {bba.id}")
        
        # Create document
        doc = Document()
        
        # Set up styles
        self._setup_styles(doc)
        
        # Set up headers and footers (different first page)
        self._setup_headers_footers(doc, bba)
        
        # Add content
        # Title page is followed immediately by the Executive Summary; we no longer
        # insert a manual page break here, otherwise Word can create an extra
        # blank page between the title page and page 2.
        self._add_title_page(doc, bba)
        
        if bba.executive_summary:
            self._add_executive_summary(doc, bba)
        
        if bba.snapshot_table:
            self._add_snapshot_table(doc, bba)
        
        if bba.expanded_findings:
            self._add_key_findings(doc, bba)
        
        if bba.twelve_month_plan:
            self._add_recommendations_plan(doc, bba)
        
        # Footer is already set up in _setup_headers_footers
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info(f"[BBA Export] Word report generated successfully")
        return buffer.getvalue()
    
    def _setup_styles(self, doc: Document):
        """Set up document styles."""
        # Modify Normal style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Heading 1 style
        h1_style = doc.styles['Heading 1']
        h1_style.font.name = 'Calibri'
        h1_style.font.size = Pt(16)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)  # Dark blue
        
        # Heading 2 style
        h2_style = doc.styles['Heading 2']
        h2_style.font.name = 'Calibri'
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
        
        # Heading 3 style
        h3_style = doc.styles['Heading 3']
        h3_style.font.name = 'Calibri'
        h3_style.font.size = Pt(12)
        h3_style.font.bold = True
    
    def _add_page_break(self, doc: Document):
        """Add a page break."""
        doc.add_page_break()
    
    def _get_main_image_path(self) -> Optional[str]:
        """Get the path to the main full-width image."""
        base_dir = Path(__file__).resolve().parents[2]
        images_dir = base_dir / "files" / "prompts" / "bba" / "images"
        img_path = images_dir / "bba_picture.png"
        
        if img_path.exists():
            return str(img_path)
        else:
            logger.warning(f"[BBA Export] Main image not found: {img_path}")
            return None
    
    def _get_grid_image_paths(self) -> list:
        """Get list of image paths for the grid row."""
        base_dir = Path(__file__).resolve().parents[2]
        images_dir = base_dir / "files" / "prompts" / "bba" / "images"
        
        # Define image filenames in the order you want them displayed
        image_files = [
            "unnamed (2).jpg",
            "unnamed (3).jpg",
            "unnamed.jpg",
            "bba_picture_3.jpg"
        ]
        
        # Build full paths and filter out non-existent files
        image_paths = []
        for img_file in image_files:
            img_path = images_dir / img_file
            if img_path.exists():
                image_paths.append(str(img_path))
            else:
                logger.warning(f"[BBA Export] Grid image not found: {img_path}")
        
        return image_paths
    
    def _add_full_width_image(self, doc: Document, image_path: str):
        """
        Add a single image at full width.
        
        Args:
            doc: Document object
            image_path: Path to the image file
        """
        try:
            # Get page width (standard US Letter is 8.5 inches, with 1 inch margins = 6.5 inches)
            # Or use section width
            section = doc.sections[0]
            page_width = section.page_width
            left_margin = section.left_margin
            right_margin = section.right_margin
            available_width = page_width - left_margin - right_margin
            
            # Create paragraph for the image
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = para.add_run()
            # Add picture with full available width
            run.add_picture(image_path, width=available_width)
            
            logger.info(f"[BBA Export] Added full-width image: {image_path}")
        except Exception as e:
            logger.error(f"[BBA Export] Failed to add full-width image {image_path}: {str(e)}")
    
    def _add_image_grid(self, doc: Document, image_paths: list, cols: int = 4, width=None):
        """
        Add images in a grid layout using a table.
        Images are same size, no gaps, full width.
        
        Args:
            doc: Document object
            image_paths: List of image file paths
            cols: Number of columns in the grid
            width: Width of each image (Inches object) - if None, calculated from full page width
        """
        if not image_paths:
            return
        
        # Get full available width (no gaps)
        section = doc.sections[0]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin
        available_width = page_width - left_margin - right_margin
        
        # Calculate image size - use full width divided by columns
        # Increase size by using 100% of available width
        if width is None:
            image_width = available_width / cols
        else:
            image_width = width
        
        # Use same height as width to ensure all images are same size
        image_height = image_width
        
        # Calculate number of rows needed
        rows = (len(image_paths) + cols - 1) // cols  # Ceiling division
        
        # Create table for grid layout
        table = doc.add_table(rows=rows, cols=cols)
        table.style = None  # No border style
        
        # Import for cell padding removal
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        # Remove all cell padding and spacing to eliminate gaps
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tbl.insert(0, tblPr)
        
        # Set table to full width (convert inches to twips: 1 inch = 1440 twips)
        # available_width is in EMU (English Metric Units), need to convert to twips
        # 1 inch = 914400 EMU, 1 inch = 1440 twips, so 1 EMU = 1440/914400 twips
        available_width_twips = int((available_width / 914400) * 1440)
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{available_width_twips}" w:type="dxa"/>')
        # Remove existing tblW if present
        for elem in tblPr:
            if elem.tag.endswith('tblW'):
                tblPr.remove(elem)
        tblPr.append(tblW)
        
        # Remove table indentation to ensure it starts at left margin
        tblInd = parse_xml(f'<w:tblInd {nsdecls("w")} w:w="0" w:type="dxa"/>')
        # Remove existing tblInd if present
        for elem in tblPr:
            if elem.tag.endswith('tblInd'):
                tblPr.remove(elem)
        tblPr.append(tblInd)
        
        # Fill table cells with images
        for idx, img_path in enumerate(image_paths):
            row_idx = idx // cols
            col_idx = idx % cols
            
            cell = table.rows[row_idx].cells[col_idx]
            cell.vertical_alignment = 1  # Center vertically
            
            # Remove cell margins/padding to eliminate gaps
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # Remove top margin
            tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}/>')
            for margin in ['top', 'left', 'bottom', 'right']:
                mar = parse_xml(f'<w:{margin} {nsdecls("w")} w:w="0" w:type="dxa"/>')
                tcMar.append(mar)
            tcPr.append(tcMar)
            
            # Remove paragraph spacing
            cell.paragraphs[0].clear()
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            try:
                run = paragraph.add_run()
                # Add picture with fixed width AND height (same size for all images)
                run.add_picture(img_path, width=image_width, height=image_height)
            except Exception as e:
                logger.error(f"[BBA Export] Failed to add image {img_path}: {str(e)}")
                # Add placeholder text if image fails
                run = paragraph.add_run(f"[Image {idx + 1}]")
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Set column widths to be equal and full width (no gaps)
        col_width = available_width / cols  # Full width divided by columns
        
        # Set width for all cells to ensure full width table
        for row in table.rows:
            for cell in row.cells:
                cell.width = col_width
    
    def _get_australia_map_path(self) -> Optional[str]:
        """Get the path to the Australia map outline image."""
        base_dir = Path(__file__).resolve().parents[2]
        images_dir = base_dir / "files" / "prompts" / "bba" / "images"
        img_path = images_dir / "bba_picture_2.png"
        
        if img_path.exists():
            return str(img_path)
        else:
            logger.warning(f"[BBA Export] Australia map image not found: {img_path}")
            return None
    
    def _add_blue_banner(self, doc: Document):
        """
        Add a blue banner with text and Australia map image.
        Full width banner with blue background.
        """
        try:
            # Get full page width (edge to edge, ignoring margins for banner)
            section = doc.sections[0]
            page_width = section.page_width
            # Use full page width for banner (not subtracting margins)
            full_page_width = page_width
            
            # Get Australia map image path
            map_image_path = self._get_australia_map_path()
            
            # Create a table with 2 columns for the banner
            table = doc.add_table(rows=1, cols=2)
            table.style = None  # No border style
            
            # Import for styling
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            
            # Set table to full page width (edge to edge)
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
                tbl.insert(0, tblPr)
            
            # Convert full page width to twips (1 inch = 1440 twips, 1 EMU = 914400 per inch)
            full_page_width_twips = int((full_page_width / 914400) * 1440)
            tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{full_page_width_twips}" w:type="dxa"/>')
            for elem in tblPr:
                if elem.tag.endswith('tblW'):
                    tblPr.remove(elem)
            tblPr.append(tblW)
            
            # Remove table indentation to ensure full width
            tblInd = parse_xml(f'<w:tblInd {nsdecls("w")} w:w="0" w:type="dxa"/>')
            for elem in tblPr:
                if elem.tag.endswith('tblInd'):
                    tblPr.remove(elem)
            tblPr.append(tblInd)
            
            # Ensure table alignment is left (no centering that might reduce width)
            for elem in tblPr:
                if elem.tag.endswith('jc'):
                    tblPr.remove(elem)
            
            # Left cell - Text
            left_cell = table.rows[0].cells[0]
            left_cell.vertical_alignment = 1  # Center vertically
            
            # Set blue background for left cell (lighter blue)
            tc = left_cell._tc
            tcPr = tc.get_or_add_tcPr()
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="5a8bb7"/>')  # Lighter blue
            # Remove existing shading if present
            for elem in tcPr:
                if elem.tag.endswith('shd'):
                    tcPr.remove(elem)
            tcPr.append(shading)
            
            # Remove cell margins
            tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}/>')
            for margin in ['top', 'left', 'bottom', 'right']:
                mar = parse_xml(f'<w:{margin} {nsdecls("w")} w:w="120" w:type="dxa"/>')  # Small padding
                tcMar.append(mar)
            # Remove existing tcMar if present
            for elem in tcPr:
                if elem.tag.endswith('tcMar'):
                    tcPr.remove(elem)
            tcPr.append(tcMar)
            
            # Add text to left cell
            left_cell.paragraphs[0].clear()
            left_para = left_cell.paragraphs[0]
            left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            left_para.paragraph_format.space_before = Pt(6)
            left_para.paragraph_format.space_after = Pt(6)
            
            text_run = left_para.add_run("Australia's Small to Medium Business Advisors")
            text_run.font.size = Pt(14)
            text_run.font.color.rgb = RGBColor(255, 255, 255)  # White text
            text_run.bold = True
            
            # Right cell - Australia map image
            right_cell = table.rows[0].cells[1]
            right_cell.vertical_alignment = 1  # Center vertically
            
            # Set blue background for right cell too (lighter blue)
            tc_right = right_cell._tc
            tcPr_right = tc_right.get_or_add_tcPr()
            shading_right = parse_xml(f'<w:shd {nsdecls("w")} w:fill="5a8bb7"/>')  # Lighter blue
            for elem in tcPr_right:
                if elem.tag.endswith('shd'):
                    tcPr_right.remove(elem)
            tcPr_right.append(shading_right)
            
            # Remove cell margins
            tcMar_right = parse_xml(f'<w:tcMar {nsdecls("w")}/>')
            for margin in ['top', 'left', 'bottom', 'right']:
                mar = parse_xml(f'<w:{margin} {nsdecls("w")} w:w="120" w:type="dxa"/>')  # Small padding
                tcMar_right.append(mar)
            for elem in tcPr_right:
                if elem.tag.endswith('tcMar'):
                    tcPr_right.remove(elem)
            tcPr_right.append(tcMar_right)
            
            # Add image to right cell
            right_cell.paragraphs[0].clear()
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right_para.paragraph_format.space_before = Pt(0)
            right_para.paragraph_format.space_after = Pt(0)
            
            if map_image_path:
                try:
                    right_run = right_para.add_run()
                    # Add map image - size it appropriately (about 1.5 inches height)
                    right_run.add_picture(map_image_path, height=Inches(1.5))
                except Exception as e:
                    logger.error(f"[BBA Export] Failed to add map image: {str(e)}")
            
            # Set column widths - make banner full page width (edge to edge)
            left_cell.width = full_page_width * 0.6  # 60% for text
            right_cell.width = full_page_width * 0.4  # 40% for image
            
            # Set cell widths explicitly in twips for both cells
            left_cell_twips = int((full_page_width * 0.6 / 914400) * 1440)
            right_cell_twips = int((full_page_width * 0.4 / 914400) * 1440)
            
            # Set cell width properties directly
            for cell in [left_cell, right_cell]:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                # Remove existing tcW if present
                for elem in tcPr:
                    if elem.tag.endswith('tcW'):
                        tcPr.remove(elem)
                # Add new tcW with full width
                if cell == left_cell:
                    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{left_cell_twips}" w:type="dxa"/>')
                else:
                    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{right_cell_twips}" w:type="dxa"/>')
                tcPr.append(tcW)
            
            logger.info(f"[BBA Export] Added blue banner with Australia map")
            
        except Exception as e:
            logger.error(f"[BBA Export] Failed to add blue banner: {str(e)}")
    
    def _add_title_page(self, doc: Document, bba: BBA):
        """Add the title page with images at the top."""
        # Add small spacing at top
        doc.add_paragraph()
        
        # Add main full-width image at the top
        main_image_path = self._get_main_image_path()
        if main_image_path:
            self._add_full_width_image(doc, main_image_path)
            doc.add_paragraph()  # Spacing after main image
        
        # Add grid of 4 images in one row
        grid_image_paths = self._get_grid_image_paths()
        if grid_image_paths:
            self._add_image_grid(doc, grid_image_paths, cols=4)
            # No spacing - banner immediately follows
        
        # Add blue banner with Australia map below the 4 images (no spacing)
        self._add_blue_banner(doc)
        doc.add_paragraph()  # Spacing after banner
        doc.add_paragraph()
        
        # If no images, add spacing
        if not main_image_path and not grid_image_paths:
            for _ in range(8):
                doc.add_paragraph()
        
        # Prepared by (moved to top with margins)
        doc.add_paragraph()  # Top margin
        prepared = doc.add_paragraph()
        prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prepared_run = prepared.add_run("Prepared by Benchmark Business Advisory")
        prepared_run.font.size = Pt(12)
        doc.add_paragraph()  # Bottom margin
        doc.add_paragraph()  # Additional spacing before title
        
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Diagnostic Findings & Recommendations Report")
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
        
        # Spacing
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Client name
        client = doc.add_paragraph()
        client.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = client.add_run(bba.client_name or "Client")
        run.font.size = Pt(18)
        
        # Date
        doc.add_paragraph()
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_str = datetime.utcnow().strftime("%B %Y")
        date_para.add_run(date_str)
    
    def _add_executive_summary(self, doc: Document, bba: BBA):
        """Add the executive summary section."""
        doc.add_heading("Executive Summary", level=1)
        
        if bba.executive_summary:
            # Split into paragraphs and add each
            paragraphs = bba.executive_summary.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text.strip())
        
        self._add_page_break(doc)
    
    def _add_snapshot_table(self, doc: Document, bba: BBA):
        """Add the Key Findings & Recommendations Snapshot table."""
        doc.add_heading("Key Findings & Recommendations Snapshot", level=1)
        
        snapshot = bba.snapshot_table
        if not snapshot:
            logger.warning(f"[BBA Export] No snapshot_table data for BBA {bba.id}")
            return
        
        # Handle both nested and non-nested data structures
        # The data might be stored as {"rows": [...]} or {"snapshot_table": {"rows": [...]}}
        if isinstance(snapshot, dict):
            if 'snapshot_table' in snapshot:
                snapshot = snapshot['snapshot_table']
            rows_data = snapshot.get('rows', [])
        else:
            logger.warning(f"[BBA Export] snapshot_table is not a dict: {type(snapshot)}")
            return
        
        if not rows_data:
            logger.warning(f"[BBA Export] No rows in snapshot_table for BBA {bba.id}")
            return
        
        logger.info(f"[BBA Export] Adding snapshot table with {len(rows_data)} rows")
        
        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row - use proper run creation for formatting
        header_cells = table.rows[0].cells
        headers = ['#', 'Priority Area', 'Key Finding', 'Recommendation']
        
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        for i, header in enumerate(headers):
            # Clear any existing content
            header_cells[i].text = ''
            # Create a new run with the header text
            run = header_cells[i].paragraphs[0].add_run(header)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
            # Set background color for header
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1a365d"/>')
            header_cells[i]._tc.get_or_add_tcPr().append(shading)
        
        # Data rows
        for row_data in rows_data:
            row = table.add_row()
            # Rank
            row.cells[0].text = str(row_data.get('rank', ''))

            # Priority Area (keep this column bold)
            priority_text = str(row_data.get('priority_area', ''))
            priority_cell = row.cells[1]
            priority_cell.text = ''
            priority_run = priority_cell.paragraphs[0].add_run(priority_text)
            priority_run.bold = True

            # Key Finding
            row.cells[2].text = str(row_data.get('key_finding', ''))

            # Recommendation
            row.cells[3].text = str(row_data.get('recommendation', ''))
        
        # Set column widths
        widths = [Inches(0.4), Inches(1.2), Inches(2.7), Inches(2.7)]
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = widths[i]
        
        logger.info(f"[BBA Export] Snapshot table added successfully")
        doc.add_paragraph()  # Spacing
        self._add_page_break(doc)
    
    def _add_key_findings(self, doc: Document, bba: BBA):
        """Add the Key Findings section."""
        doc.add_heading("Key Findings – Ranked by Importance", level=1)
        
        findings = bba.expanded_findings
        if not findings:
            return
        
        findings_list = findings.get('expanded_findings', [])
        
        for finding in findings_list:
            # Finding heading
            rank = finding.get('rank', '')
            title = finding.get('title', '')
            doc.add_heading(f"{rank}. {title}", level=2)
            
            # Priority area
            priority = finding.get('priority_area', '')
            if priority:
                para = doc.add_paragraph()
                run = para.add_run(f"Priority Area: {priority}")
                run.italic = True
            
            # Paragraphs
            paragraphs = finding.get('paragraphs', [])
            for para_text in paragraphs:
                if para_text:
                    doc.add_paragraph(para_text)
            
            # Key points
            key_points = finding.get('key_points', [])
            if key_points:
                doc.add_paragraph()
                for point in key_points:
                    doc.add_paragraph(point, style='List Bullet')
            
            doc.add_paragraph()  # Spacing between findings
        
        self._add_page_break(doc)
    
    def _add_recommendations_plan(self, doc: Document, bba: BBA):
        """Add the 12-Month Recommendations Plan section."""
        doc.add_heading("12-Month Recommendations Plan", level=1)
        
        plan = bba.twelve_month_plan
        if not plan:
            return
        
        # Plan notes
        plan_notes = plan.get('plan_notes', '') or bba.plan_notes
        if plan_notes:
            doc.add_heading("Notes on the 12-Month Recommendations Plan", level=2)
            doc.add_paragraph(plan_notes)
            doc.add_paragraph()
        
        # Recommendations
        recommendations = plan.get('recommendations', [])
        
        for rec in recommendations:
            # Recommendation heading
            number = rec.get('number', '')
            title = rec.get('title', '')
            timing = rec.get('timing', '')
            doc.add_heading(f"Recommendation {number}: {title}", level=2)
            
            # Timing
            if timing:
                para = doc.add_paragraph()
                run = para.add_run(f"Timing: {timing}")
                run.italic = True
            
            # Purpose
            purpose = rec.get('purpose', '')
            if purpose:
                doc.add_heading("Purpose", level=3)
                doc.add_paragraph(purpose)
            
            # Key Objectives
            objectives = rec.get('key_objectives', [])
            if objectives:
                doc.add_heading("Key Objectives", level=3)
                for obj in objectives:
                    doc.add_paragraph(obj, style='List Bullet')
            
            # Actions to Complete
            actions = rec.get('actions', [])
            if actions:
                doc.add_heading("Actions to Complete", level=3)
                for i, action in enumerate(actions, 1):
                    doc.add_paragraph(f"{i}. {action}")
            
            # BBA Support Outline
            bba_support = rec.get('bba_support', '')
            if bba_support:
                doc.add_heading("BBA Support Outline", level=3)
                doc.add_paragraph(bba_support)
            
            # Expected Outcomes
            outcomes = rec.get('expected_outcomes', [])
            if outcomes:
                doc.add_heading("Expected Outcomes", level=3)
                for outcome in outcomes:
                    doc.add_paragraph(outcome, style='List Bullet')
            
            doc.add_paragraph()  # Spacing between recommendations
        
        # Timeline Summary
        timeline = plan.get('timeline_summary', {})
        if timeline and timeline.get('rows'):
            self._add_page_break(doc)
            doc.add_heading("Implementation Timeline", level=2)
            
            rows_data = timeline.get('rows', [])
            
            # Create table
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            headers = ['Rec #', 'Recommendation', 'Focus Area', 'Timing', 'Key Outcome']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # Data rows
            for row_data in rows_data:
                row = table.add_row()
                row.cells[0].text = str(row_data.get('rec_number', ''))
                row.cells[1].text = row_data.get('recommendation', '')
                row.cells[2].text = row_data.get('focus_area', '')
                row.cells[3].text = row_data.get('timing', '')
                row.cells[4].text = row_data.get('key_outcome', '')
    
    def _get_logo_path(self) -> Optional[str]:
        """Get the path to the Benchmark logo."""
        base_dir = Path(__file__).resolve().parents[2]
        # Try multiple possible locations
        possible_paths = [
            base_dir / "files" / "prompts" / "bba" / "images" / "logo.png",
            base_dir / "files" / "prompts" / "bba" / "images" / "benchmark_logo.png",
            base_dir / "files" / "prompts" / "bba" / "images" / "unnamed (1).png",
        ]
        
        for img_path in possible_paths:
            if img_path.exists():
                return str(img_path)
        
        logger.warning(f"[BBA Export] Logo not found, tried: {possible_paths}")
        return None
    
    def _setup_headers_footers(self, doc: Document, bba: BBA):
        """Set up headers and footers for all pages (except first page)."""
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        section = doc.sections[0]
        
        # Enable different first page header/footer
        sectPr = section._sectPr
        if sectPr is None:
            sectPr = parse_xml(f'<w:sectPr {nsdecls("w")}/>')
            section._element.append(sectPr)
        
        # Set different first page
        titlePg = parse_xml(f'<w:titlePg {nsdecls("w")}/>')
        # Remove existing titlePg if present
        for elem in sectPr:
            if elem.tag.endswith('titlePg'):
                sectPr.remove(elem)
        sectPr.append(titlePg)
        
        # Setup header (for pages 2+)
        header = section.header
        header.is_linked_to_previous = False
        
        # Clear existing paragraphs
        for para in header.paragraphs:
            para.clear()
        
        # Calculate section width first
        section_width = section.page_width - section.left_margin - section.right_margin
        
        # Create header table with logo on right and text on left
        header_table = header.add_table(rows=1, cols=2, width=section_width)
        header_table.style = None
        
        # Set table to full width (additional XML configuration for compatibility)
        tbl = header_table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tbl.insert(0, tblPr)
        
        available_width_twips = int((section_width / 914400) * 1440)
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{available_width_twips}" w:type="dxa"/>')
        for elem in tblPr:
            if elem.tag.endswith('tblW'):
                tblPr.remove(elem)
        tblPr.append(tblW)
        
        # Left cell - Logo (moved from right to left, smaller size)
        left_cell = header_table.rows[0].cells[0]
        left_cell.vertical_alignment = 1
        left_para = left_cell.paragraphs[0]
        left_para.clear()
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        logo_path = self._get_logo_path()
        if logo_path:
            try:
                logo_run = left_para.add_run()
                # Make logo smaller - reduced from 0.5 to 0.3 inches
                logo_run.add_picture(logo_path, height=Inches(0.3))
            except Exception as e:
                logger.error(f"[BBA Export] Failed to add logo to header: {str(e)}")
        
        # Right cell - Confidential text (moved from left to right)
        right_cell = header_table.rows[0].cells[1]
        right_cell.vertical_alignment = 1
        right_para = right_cell.paragraphs[0]
        right_para.clear()
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Get date with ordinal suffix (e.g., "20th Oct 2025")
        date_obj = datetime.utcnow()
        day = date_obj.day
        # Add ordinal suffix (st, nd, rd, th)
        if 4 <= day <= 20 or 24 <= day <= 30:
            suffix = "th"
        else:
            suffix = ["st", "nd", "rd"][day % 10 - 1]
        
        # Format: "20th Oct 2025"
        month_abbr = date_obj.strftime("%b")
        date_str = f"{day}{suffix} {month_abbr} {date_obj.year}"
        
        conf_run = right_para.add_run("CONFIDENTIAL - SAMPLE REPORT – ")
        conf_run.font.size = Pt(10)
        conf_run.font.color.rgb = RGBColor(64, 64, 64)
        
        date_run = right_para.add_run(date_str)
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(64, 64, 64)
        
        # Set column widths - logo on left takes less space, text on right takes more
        left_cell.width = section_width * 0.2  # 20% for logo
        right_cell.width = section_width * 0.8  # 80% for text
        
        # Setup footer (for pages 2+)
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # Clear existing paragraphs
        for para in footer.paragraphs:
            para.clear()
        
        # We use a 3‑column table so the page number can sit on the bottom‑right,
        footer_table = footer.add_table(rows=1, cols=3, width=section_width)
        footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # fits on a single physical line.
        tbl = footer_table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tbl.insert(0, tblPr)
        
        # Table width in twips
        available_width_twips = int((section_width / 914400) * 1440)
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{available_width_twips}" w:type="dxa"/>')
        for elem in tblPr:
            if elem.tag.endswith('tblW'):
                tblPr.remove(elem)
        tblPr.append(tblW)
        footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set relative column widths: left 20%, middle 40%, right 40%
        # (this visually centers the email while keeping enough room on the right
        #  so "1300 366 521 | Page X of Y" stays on one line)
        width_ratios = [0.20, 0.40, 0.40]
        row = footer_table.rows[0]
        for idx, cell in enumerate(row.cells):
            cell_twips = int(available_width_twips * width_ratios[idx])
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for elem in tcPr:
                if elem.tag.endswith('tcW'):
                    tcPr.remove(elem)
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{cell_twips}" w:type="dxa"/>')
            tcPr.append(tcW)
        
        left_cell, middle_cell, right_cell = footer_table.rows[0].cells
        left_para = left_cell.paragraphs[0]
        middle_para = middle_cell.paragraphs[0]
        right_para = right_cell.paragraphs[0]
        
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        middle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add footer text with hyperlink for email
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        # Left: website (slightly smaller to help everything fit on one line)
        website_run = left_para.add_run("BenchmarkBusinessAdvisory.com.au")
        website_run.font.size = Pt(8)
        website_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Middle: email as hyperlink
        email_text = "chat@benchmarkbusinessadvisory.com.au"
        email_url = f"mailto:{email_text}"
        
        try:
            # Get the document part to add relationship
            part = footer.part
            # Add hyperlink relationship
            r_id = part.relate_to(
                email_url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            )
            
            # Create hyperlink element with run inside (explicit small font size)
            hyperlink_elem = parse_xml(
                f'<w:hyperlink {nsdecls("w", "r")} r:id="{r_id}">'
                f'<w:r><w:rPr><w:color w:val="00007A"/><w:u w:val="single"/><w:sz w:val="16"/></w:rPr>'
                f'<w:t>{email_text}</w:t></w:r></w:hyperlink>'
            )
            
            # Append hyperlink to paragraph
            middle_para._p.append(hyperlink_elem)
            
        except Exception as e:
            logger.warning(f"[BBA Export] Failed to create email hyperlink, using styled text: {str(e)}")
            # Fallback: add email as styled text (looks like a link but not clickable)
            email_run = middle_para.add_run(email_text)
            email_run.font.size = Pt(8)
            email_run.font.color.rgb = RGBColor(0, 0, 122)  # Blue
            email_run.underline = True
        
        # Right: phone + "Page X of Y" (true bottom‑right page count)
        phone_run = right_para.add_run("1300 366 521 | Page ")
        phone_run.font.size = Pt(8)
        phone_run.font.color.rgb = RGBColor(0, 0, 0)
        
        def _add_field_run(paragraph, instruction: str):
            """Insert a Word field (e.g. PAGE, NUMPAGES) into a paragraph."""
            field_run = paragraph.add_run()
            field_run.font.size = Pt(8)
            field_run.font.color.rgb = RGBColor(0, 0, 0)
            
            run_element = field_run._element
            run_element.clear()
            
            rPr = parse_xml(
                f'<w:rPr {nsdecls("w")}><w:color w:val="000000"/><w:sz w:val="16"/></w:rPr>'
            )
            run_element.append(rPr)
            
            fld_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            instr_text = parse_xml(
                f'<w:instrText {nsdecls("w")}>{instruction}</w:instrText>'
            )
            fld_separate = parse_xml(
                f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'
            )
            text_elem = parse_xml(f'<w:t {nsdecls("w")}/>')
            fld_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            
            for el in (fld_begin, instr_text, fld_separate, text_elem, fld_end):
                run_element.append(el)
        
        # Current page number
        _add_field_run(right_para, "PAGE")
        
        # " of " total pages
        of_run = right_para.add_run(" of ")
        of_run.font.size = Pt(8)
        of_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Total page count
        _add_field_run(right_para, "NUMPAGES")
    
    def _add_footer(self, doc: Document, bba: BBA):
        """Add footer to document (legacy method - now handled by _setup_headers_footers)."""
        # This method is kept for backward compatibility but now uses the new setup
        self._setup_headers_footers(doc, bba)


# Factory function
def get_bba_report_exporter() -> BBAReportExporter:
    """Get a BBA report exporter instance."""
    return BBAReportExporter()
