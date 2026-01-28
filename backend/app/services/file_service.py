"""
File upload and management service
"""
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session
from typing import List, Optional, Tuple
from uuid import UUID
import uuid
import os
import shutil
from pathlib import Path
import tempfile
import logging

from app.models.media import Media
from app.models.user import User
from app.services.openai_service import openai_service
from app.config import settings

logger = logging.getLogger(__name__)


class FileService:
    """Service for handling file uploads and management"""
    
    # Allowed file extensions for diagnostics
    ALLOWED_EXTENSIONS = {
        # Documents
        'pdf', 'doc', 'docx', 'txt', 'rtf',
        # Spreadsheets
        'xls', 'xlsx', 'csv',
        # Images
        'jpg', 'jpeg', 'png', 'gif', 'webp',
        # Other
        'zip'
    }
    
    # Max file size: 10MB
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    def __init__(self, db: Session):
        self.db = db
        # Use the singleton openai_service instance
        self.openai_service = openai_service
        # Use files/uploads as the base upload directory
        # Path(__file__) = backend/app/services/file_service.py
        # .parents[2] = backend/
        # / "files" / "uploads" = backend/files/uploads
        base_dir = Path(__file__).resolve().parents[2]  # Go up to backend/
        self.upload_dir = base_dir / "files" / "uploads"
        self.upload_dir.mkdir(parents=True, exist_ok=True)
    
    def _get_file_extension(self, filename: str) -> str:
        """Extract file extension from filename"""
        return filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    
    def _validate_file(self, file: UploadFile) -> None:
        """Validate uploaded file"""
        # Check extension
        ext = self._get_file_extension(file.filename)
        if ext not in self.ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"File type .{ext} not allowed. Allowed types: {', '.join(self.ALLOWED_EXTENSIONS)}"
            )
        
        # Check file size (if available)
        if hasattr(file, 'size') and file.size and file.size > self.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"File size exceeds maximum allowed size of {self.MAX_FILE_SIZE / (1024*1024)}MB"
            )
    
    def _should_convert_to_pdf(self, file_extension: str) -> bool:
        """
        Determine if a file should be converted to PDF before sending to OpenAI
        
        Args:
            file_extension: File extension (e.g., 'docx', 'xlsx')
            
        Returns:
            True if file should be converted to PDF
        """
        # Convert these formats to PDF to prevent code interpreter processing
        convertible_formats = {
            'doc', 'docx', 'txt', 'rtf',  # Documents
            'xls', 'xlsx', 'csv',  # Spreadsheets
            'jpg', 'jpeg', 'png', 'gif', 'webp'  # Images
        }
        return file_extension.lower() in convertible_formats
    
    def _convert_to_pdf(self, input_path: Path) -> Optional[Path]:
        """
        Convert a file to PDF format (optimized for speed)
        
        Args:
            input_path: Path to the input file
            
        Returns:
            Path to the converted PDF file, or None if conversion failed
        """
        try:
            file_ext = input_path.suffix.lower().lstrip('.')
            output_path = input_path.parent / f"{input_path.stem}_converted.pdf"
            
            logger.info(f"⚡ Converting {file_ext} file to PDF: {input_path.name}")
            
            # Convert based on file type
            if file_ext in ['doc', 'docx']:
                return self._convert_docx_to_pdf(input_path, output_path)
            elif file_ext in ['xls', 'xlsx']:
                return self._convert_xlsx_to_pdf(input_path, output_path)
            elif file_ext == 'csv':
                return self._convert_csv_to_pdf(input_path, output_path)
            elif file_ext in ['jpg', 'jpeg', 'png', 'gif', 'webp']:
                return self._convert_image_to_pdf(input_path, output_path)
            elif file_ext in ['txt', 'rtf']:
                return self._convert_text_to_pdf(input_path, output_path)
            else:
                logger.warning(f"No converter available for {file_ext}")
                return None
                
        except Exception as e:
            logger.error(f"Failed to convert {input_path.name} to PDF: {str(e)}", exc_info=True)
            return None
    
    def _convert_docx_to_pdf(self, input_path: Path, output_path: Path) -> Optional[Path]:
        """Convert DOCX/DOC to PDF"""
        try:
            # Try using docx2pdf (works best on Windows)
            try:
                from docx2pdf import convert
                convert(str(input_path), str(output_path))
                logger.info(f"Successfully converted DOCX to PDF using docx2pdf")
                return output_path
            except ImportError:
                logger.warning("docx2pdf not available, falling back to python-docx + reportlab")
            
            # Fallback: Use python-docx + reportlab
            from docx import Document
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table as PDFTable, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib import colors
            
            doc = Document(str(input_path))
            pdf = SimpleDocTemplate(str(output_path), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            logger.info(f"Processing DOCX: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
            
            # Process all paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    p = Paragraph(para.text, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 12))
            
            # Process all tables
            for table_idx, table in enumerate(doc.tables):
                logger.info(f"Processing table {table_idx + 1}/{len(doc.tables)}: {len(table.rows)} rows × {len(table.columns)} columns")
                
                # Extract table data
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    # Create PDF table
                    pdf_table = PDFTable(table_data, repeatRows=1)
                    pdf_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 8),
                        ('FONTSIZE', (0, 1), (-1, -1), 7),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ]))
                    story.append(Spacer(1, 12))
                    story.append(pdf_table)
                    story.append(Spacer(1, 12))
            
            pdf.build(story)
            logger.info(f"✅ Successfully converted DOCX to PDF using python-docx + reportlab (FULL FILE)")
            return output_path
            
        except Exception as e:
            logger.error(f"DOCX to PDF conversion failed: {str(e)}")
            return None
    
    def _convert_xlsx_to_pdf(self, input_path: Path, output_path: Path) -> Optional[Path]:
        """Convert XLSX/XLS to PDF - Full file conversion (all sheets, all rows, all columns)"""
        try:
            import pandas as pd
            from reportlab.lib.pagesizes import letter, landscape
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, PageBreak, Paragraph, Spacer
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet
            
            logger.info(f"Starting Excel to PDF conversion (FULL FILE - all data)...")
            
            # Read Excel file - process ALL sheets
            excel_file = pd.ExcelFile(input_path, engine='openpyxl')
            pdf = SimpleDocTemplate(str(output_path), pagesize=landscape(letter))
            styles = getSampleStyleSheet()
            story = []
            
            # Process ALL sheets (no limit)
            total_sheets = len(excel_file.sheet_names)
            logger.info(f"Processing {total_sheets} sheet(s): {', '.join(excel_file.sheet_names)}")
            
            for idx, sheet_name in enumerate(excel_file.sheet_names):
                logger.info(f"Processing sheet {idx + 1}/{total_sheets}: {sheet_name}")
                
                # Read FULL sheet (no row limit)
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                logger.info(f"  Sheet '{sheet_name}': {len(df)} rows × {len(df.columns)} columns")
                
                # Add sheet title
                story.append(Paragraph(f"<b>{sheet_name}</b>", styles['Heading1']))
                story.append(Spacer(1, 12))
                
                # Process ALL columns (no limit)
                # Prepare table data - truncate cell content if too long to prevent PDF errors
                max_cell_length = 200  # Limit individual cell text length
                df_display = df.fillna('').astype(str)
                # Truncate long cell values - compatible with all pandas versions
                def truncate_cell(x):
                    x_str = str(x)
                    return x_str[:max_cell_length] + '...' if len(x_str) > max_cell_length else x_str
                
                # Use map() for pandas 2.1+, applymap() for older versions
                import pandas as pd
                if pd.__version__ >= '2.1.0':
                    df_display = df_display.map(truncate_cell)
                else:
                    df_display = df_display.applymap(truncate_cell)
                
                data = [df_display.columns.tolist()] + df_display.values.tolist()
                
                # Create table - use smaller font to fit more data
                table = Table(data, repeatRows=1)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 7),
                    ('FONTSIZE', (0, 1), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                
                story.append(table)
                story.append(PageBreak())
            
            logger.info(f"Building PDF document with full data...")
            pdf.build(story)
            logger.info(f"✅ Successfully converted XLSX to PDF (FULL FILE - {total_sheets} sheet(s))")
            return output_path
            
        except Exception as e:
            logger.error(f"XLSX to PDF conversion failed: {str(e)}", exc_info=True)
            return None
    
    def _convert_csv_to_pdf(self, input_path: Path, output_path: Path) -> Optional[Path]:
        """Convert CSV to PDF - Full file conversion (all rows, all columns)"""
        try:
            import pandas as pd
            from reportlab.lib.pagesizes import letter, landscape
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet
            
            logger.info(f"Starting CSV to PDF conversion (FULL FILE - all data)...")
            
            # Read FULL CSV file (no row limit)
            df = pd.read_csv(input_path)
            
            logger.info(f"CSV file: {len(df)} rows × {len(df.columns)} columns")
            
            pdf = SimpleDocTemplate(str(output_path), pagesize=landscape(letter))
            styles = getSampleStyleSheet()
            
            # Process ALL columns (no limit)
            # Truncate cell content if too long to prevent PDF errors
            max_cell_length = 200  # Limit individual cell text length
            df_display = df.fillna('').astype(str)
            # Truncate long cell values - compatible with all pandas versions
            def truncate_cell(x):
                x_str = str(x)
                return x_str[:max_cell_length] + '...' if len(x_str) > max_cell_length else x_str
            
            # Use map() for pandas 2.1+, applymap() for older versions
            import pandas as pd
            if pd.__version__ >= '2.1.0':
                df_display = df_display.map(truncate_cell)
            else:
                df_display = df_display.applymap(truncate_cell)
            
            # Prepare table data
            data = [df_display.columns.tolist()] + df_display.values.tolist()
            
            # Create table - use smaller font to fit more data
            table = Table(data, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 7),
                ('FONTSIZE', (0, 1), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            
            story = [table]
            
            logger.info(f"Building PDF document with full data...")
            pdf.build(story)
            logger.info(f"✅ Successfully converted CSV to PDF (FULL FILE - {len(df)} rows)")
            return output_path
            
        except Exception as e:
            logger.error(f"CSV to PDF conversion failed: {str(e)}", exc_info=True)
            return None
    
    def _convert_image_to_pdf(self, input_path: Path, output_path: Path) -> Optional[Path]:
        """Convert image to PDF"""
        try:
            from PIL import Image
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            
            # Open image
            img = Image.open(input_path)
            
            # Convert RGBA to RGB if necessary
            if img.mode == 'RGBA':
                img = img.convert('RGB')
            
            # Create PDF
            c = canvas.Canvas(str(output_path), pagesize=letter)
            
            # Get image dimensions
            img_width, img_height = img.size
            
            # Scale image to fit page while maintaining aspect ratio
            page_width, page_height = letter
            margin = 50
            available_width = page_width - 2 * margin
            available_height = page_height - 2 * margin
            
            # Calculate scaling
            width_scale = available_width / img_width
            height_scale = available_height / img_height
            scale = min(width_scale, height_scale)
            
            final_width = img_width * scale
            final_height = img_height * scale
            
            # Center image on page
            x = (page_width - final_width) / 2
            y = (page_height - final_height) / 2
            
            # Save image temporarily as it needs to be in a format reportlab can handle
            temp_img_path = input_path.parent / f"temp_{input_path.stem}.jpg"
            img.save(temp_img_path, 'JPEG')
            
            # Draw image on canvas
            c.drawImage(str(temp_img_path), x, y, width=final_width, height=final_height)
            c.save()
            
            # Clean up temp image
            if temp_img_path.exists():
                temp_img_path.unlink()
            
            logger.info(f"Successfully converted image to PDF")
            return output_path
            
        except Exception as e:
            logger.error(f"Image to PDF conversion failed: {str(e)}")
            return None
    
    def _convert_text_to_pdf(self, input_path: Path, output_path: Path) -> Optional[Path]:
        """Convert text file to PDF - Full file conversion"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch
            
            logger.info(f"Starting text to PDF conversion (FULL FILE - all content)...")
            
            # Read FULL text file
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()
            
            file_size = len(text_content)
            logger.info(f"Text file: {file_size:,} characters")
            
            pdf = SimpleDocTemplate(str(output_path), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Split into paragraphs - process ALL content
            paragraphs = text_content.split('\n')
            logger.info(f"Processing {len(paragraphs)} lines")
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Escape special characters for reportlab
                    para_text = para_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    p = Paragraph(para_text, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 0.2 * inch))
            
            pdf.build(story)
            logger.info(f"✅ Successfully converted text to PDF (FULL FILE - {file_size:,} characters)")
            return output_path
            
        except Exception as e:
            logger.error(f"Text to PDF conversion failed: {str(e)}", exc_info=True)
            return None
    
    async def upload_file(
        self,
        file: UploadFile,
        user_id: UUID,
        question_field_name: Optional[str] = None,
        description: Optional[str] = None,
        upload_to_openai: bool = True,
        diagnostic_id: Optional[UUID] = None
    ) -> Media:
        """
        Upload a file, store it locally, and optionally upload to OpenAI
        
        Args:
            file: The uploaded file
            user_id: ID of the user uploading the file
            question_field_name: Which diagnostic question this file answers
            description: Optional description
            upload_to_openai: Whether to upload to OpenAI for analysis
            diagnostic_id: Optional diagnostic ID. If provided, files are stored in 
                          files/uploads/diagnostic/{diagnostic_id}/, otherwise files/uploads/users/{user_id}/
            
        Returns:
            Media object representing the uploaded file
        """
        # Validate file
        self._validate_file(file)
        
        # Create directory based on whether diagnostic_id is provided
        if diagnostic_id:
            # Store diagnostic files in files/uploads/diagnostic/{diagnostic_id}/
            storage_dir = self.upload_dir / "diagnostic" / str(diagnostic_id)
        else:
            # Store user files (profile pictures, etc.) in files/uploads/users/{user_id}/
            storage_dir = self.upload_dir / "users" / str(user_id)
        
        storage_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate unique filename
        ext = self._get_file_extension(file.filename)
        unique_filename = f"{uuid.uuid4()}.{ext}"
        file_path = storage_dir / unique_filename
        
        # Save file locally
        try:
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")
        
        # Get file size
        file_size = os.path.getsize(file_path)
        
        # Create media record
        media = Media(
            user_id=user_id,
            file_name=file.filename,
            file_path=str(file_path),
            file_size=file_size,
            file_type=file.content_type,
            file_extension=ext,
            question_field_name=question_field_name,
            description=description
        )
        
        self.db.add(media)
        self.db.flush()  # Get the media ID
        
        # Upload to OpenAI if requested
        if upload_to_openai:
            pdf_path_for_cleanup = None
            try:
                # Determine which file to upload to OpenAI
                file_to_upload = file_path
                
                # Convert to PDF if it's a convertible format
                was_converted_to_pdf = False
                if self._should_convert_to_pdf(ext):
                    logger.info(f"File {file.filename} will be converted to PDF before OpenAI upload")
                    converted_pdf = self._convert_to_pdf(file_path)
                    
                    if converted_pdf and converted_pdf.exists():
                        file_to_upload = converted_pdf
                        pdf_path_for_cleanup = converted_pdf
                        was_converted_to_pdf = True
                        logger.info(f"✅ Using converted PDF for OpenAI upload: {converted_pdf.name}")
                        
                        # CRITICAL: Update Media record to reflect PDF extension
                        # This ensures diagnostic_service categorizes it as PDF, not Code Interpreter
                        media.file_extension = "pdf"
                        media.file_type = "application/pdf"
                        logger.info(f"📝 Updated Media record: file_extension=pdf, file_type=application/pdf")
                    else:
                        logger.warning(f"PDF conversion failed, uploading original file")
                else:
                    logger.info(f"File {file.filename} is already PDF or doesn't need conversion")
                
                file_path_str = str(file_to_upload)
                logger.info(f"📤 Uploading file to OpenAI from path: {file_path_str}")
                openai_file = await self.openai_service.upload_file(
                    file_path=file_path_str,
                    purpose="user_data"  # For Responses API + tools (e.g., code_interpreter)
                )
                
                if openai_file:
                    media.openai_file_id = openai_file.get('id')
                    media.openai_purpose = openai_file.get('purpose') or "user_data"
                    from datetime import datetime
                    media.openai_uploaded_at = datetime.utcnow()
                    if was_converted_to_pdf:
                        logger.info(f"✅ PDF file uploaded to OpenAI: {media.openai_file_id} (converted from {ext})")
                    else:
                        logger.info(f"✅ File uploaded to OpenAI: {media.openai_file_id}")
                    
                # Clean up converted PDF if it was created
                if pdf_path_for_cleanup and pdf_path_for_cleanup.exists():
                    try:
                        pdf_path_for_cleanup.unlink()
                        logger.info(f"Cleaned up converted PDF: {pdf_path_for_cleanup.name}")
                    except Exception as cleanup_error:
                        logger.warning(f"Failed to cleanup converted PDF: {str(cleanup_error)}")
                        
            except Exception as e:
                logger.error(f"Failed to upload file to OpenAI: {str(e)}")
                # Clean up converted PDF on error
                if pdf_path_for_cleanup and pdf_path_for_cleanup.exists():
                    try:
                        pdf_path_for_cleanup.unlink()
                    except:
                        pass
                # Continue even if OpenAI upload fails
        
        self.db.commit()
        self.db.refresh(media)
        
        return media
    
    async def upload_files(
        self,
        files: List[UploadFile],
        user_id: UUID,
        question_field_name: Optional[str] = None,
        upload_to_openai: bool = True,
        diagnostic_id: Optional[UUID] = None
    ) -> List[Media]:
        """
        Upload multiple files
        
        Args:
            files: List of uploaded files
            user_id: ID of the user uploading the files
            question_field_name: Which diagnostic question these files answer
            upload_to_openai: Whether to upload to OpenAI
            diagnostic_id: Optional diagnostic ID. If provided, files are stored in 
                          files/uploads/diagnostic/{diagnostic_id}/, otherwise files/uploads/users/{user_id}/
            
        Returns:
            List of Media objects
        """
        media_list = []
        
        for file in files:
            try:
                media = await self.upload_file(
                    file=file,
                    user_id=user_id,
                    question_field_name=question_field_name,
                    upload_to_openai=upload_to_openai,
                    diagnostic_id=diagnostic_id
                )
                media_list.append(media)
            except Exception as e:
                print(f"  Failed to upload file {file.filename}: {str(e)}")
                # Continue with other files
        
        return media_list
    
    def get_user_files(self, user_id: UUID) -> List[Media]:
        """Get all files for a user"""
        return self.db.query(Media).filter(
            Media.user_id == user_id,
            Media.is_active == True
        ).order_by(Media.created_at.desc()).all()
    
    def get_diagnostic_files(self, diagnostic_id: UUID) -> List[Media]:
        """Get all files attached to a diagnostic"""
        from app.models.diagnostic import Diagnostic
        
        diagnostic = self.db.query(Diagnostic).filter(
            Diagnostic.id == diagnostic_id
        ).first()
        
        if not diagnostic:
            return []
        
        return list(diagnostic.media)
    
    def attach_file_to_diagnostic(
        self,
        media_id: UUID,
        diagnostic_id: UUID
    ) -> bool:
        """Attach a file to a diagnostic"""
        from app.models.diagnostic import Diagnostic
        
        media = self.db.query(Media).filter(Media.id == media_id).first()
        diagnostic = self.db.query(Diagnostic).filter(Diagnostic.id == diagnostic_id).first()
        
        if not media or not diagnostic:
            return False
        
        if media not in diagnostic.media:
            diagnostic.media.append(media)
            self.db.commit()
        
        return True
    
    def delete_file(self, media_id: UUID, hard_delete: bool = False) -> bool:
        """
        Delete a file (soft delete by default)
        
        Args:
            media_id: ID of the media to delete
            hard_delete: If True, permanently delete; if False, soft delete
            
        Returns:
            True if successful
        """
        media = self.db.query(Media).filter(Media.id == media_id).first()
        
        if not media:
            return False
        
        if hard_delete:
            # Delete physical file
            try:
                if os.path.exists(media.file_path):
                    os.remove(media.file_path)
            except Exception as e:
                print(f"  Failed to delete physical file: {str(e)}")
            
            # Delete from database
            self.db.delete(media)
        else:
            # Soft delete
            from datetime import datetime
            media.is_active = False
            media.deleted_at = datetime.utcnow()
        
        self.db.commit()
        return True


    """Dependency injection for FileService"""
def get_file_service(db: Session) -> FileService:
    return FileService(db)

