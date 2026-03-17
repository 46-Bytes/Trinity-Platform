"""
Document Template Service
Generates documents from Word templates stored in the database
"""
from typing import List, Dict, Any, Optional
from uuid import UUID
import io
import re
import logging

from sqlalchemy.orm import Session

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.models.document_template import DocumentTemplate

logger = logging.getLogger(__name__)


class DocumentTemplateService:
    """
    Service for generating documents from Word templates stored in the database.
    """

    PLACEHOLDER_PATTERN = re.compile(r'\{\{(\w+)\}\}')

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for document template generation. "
                "Install it with: pip install python-docx"
            )

    def list_available_templates(self, db: Session) -> List[Dict[str, str]]:
        """List all available templates from the database."""
        templates = db.query(DocumentTemplate).order_by(DocumentTemplate.display_name).all()
        return [
            {"name": t.file_name, "display_name": t.display_name}
            for t in templates
        ]

    def get_template(self, db: Session, file_name: str) -> Optional[DocumentTemplate]:
        """Get a template by filename."""
        return db.query(DocumentTemplate).filter(DocumentTemplate.file_name == file_name).first()

    def upload_template(self, db: Session, file_name: str, file_data: bytes, uploaded_by_user_id: UUID) -> DocumentTemplate:
        """Upload a new template to the database."""
        # Generate display name from filename
        stem = file_name.rsplit('.', 1)[0] if '.' in file_name else file_name
        display_name = stem.replace("_", " ").replace("-", " ")
        display_name = " ".join(word.capitalize() for word in display_name.split())

        template = DocumentTemplate(
            file_name=file_name,
            display_name=display_name,
            file_data=file_data,
            file_size=len(file_data),
            uploaded_by_user_id=uploaded_by_user_id,
        )
        db.add(template)
        db.commit()
        db.refresh(template)
        logger.info(f"Template uploaded to database: {file_name} by user {uploaded_by_user_id}")
        return template

    def delete_template(self, db: Session, file_name: str) -> bool:
        """Delete a template from the database. Returns True if deleted."""
        template = self.get_template(db, file_name)
        if not template:
            return False
        db.delete(template)
        db.commit()
        logger.info(f"Template deleted from database: {file_name}")
        return True

    def extract_placeholders_from_bytes(self, file_data: bytes) -> List[str]:
        """Extract all placeholders from template file data."""
        placeholders = set()
        doc = Document(io.BytesIO(file_data))

        for paragraph in doc.paragraphs:
            matches = self.PLACEHOLDER_PATTERN.findall(paragraph.text)
            placeholders.update(matches)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = self.PLACEHOLDER_PATTERN.findall(paragraph.text)
                        placeholders.update(matches)

        return sorted(list(placeholders))

    def generate_document(
        self,
        db: Session,
        template_name: str,
        user_responses: Dict[str, Any]
    ) -> bytes:
        """
        Generate a document from a database-stored template by replacing placeholders.

        Returns:
            Bytes of the generated Word document
        """
        template = self.get_template(db, template_name)
        if not template:
            raise FileNotFoundError(f"Template not found: {template_name}")

        logger.info(f"Generating document from template: {template_name}")

        try:
            doc = Document(io.BytesIO(template.file_data))

            # Extract placeholders
            placeholders = self.extract_placeholders_from_bytes(template.file_data)
            logger.info(f"Found {len(placeholders)} unique placeholders in template")

            # Build replacements dictionary
            replacements = {}
            missing_fields = []

            for placeholder in placeholders:
                if placeholder in user_responses:
                    value = user_responses[placeholder]
                    if value is None:
                        replacements[placeholder] = "[Not provided]"
                    elif isinstance(value, (list, dict)):
                        replacements[placeholder] = str(value)
                    else:
                        replacements[placeholder] = str(value)
                else:
                    missing_fields.append(placeholder)
                    replacements[placeholder] = "[Not provided]"

            if missing_fields:
                logger.warning(f"Missing fields in user_responses: {missing_fields}")

            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text = paragraph.text
                    for placeholder, replacement in replacements.items():
                        text = text.replace(f"{{{{{placeholder}}}}}", replacement)

                    if text != paragraph.text:
                        if paragraph.runs:
                            first_run = paragraph.runs[0]
                            paragraph.clear()
                            new_run = paragraph.add_run(text)
                            new_run.font.name = first_run.font.name
                            try:
                                new_run.font.size = first_run.font.size
                            except (ValueError, TypeError):
                                pass
                            new_run.font.bold = first_run.font.bold
                            new_run.font.italic = first_run.font.italic
                        else:
                            paragraph.text = text

            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text:
                                text = paragraph.text
                                for placeholder, replacement in replacements.items():
                                    text = text.replace(f"{{{{{placeholder}}}}}", replacement)

                                if text != paragraph.text:
                                    if paragraph.runs:
                                        first_run = paragraph.runs[0]
                                        paragraph.clear()
                                        new_run = paragraph.add_run(text)
                                        new_run.font.name = first_run.font.name
                                        try:
                                            new_run.font.size = first_run.font.size
                                        except (ValueError, TypeError):
                                            pass
                                        new_run.font.bold = first_run.font.bold
                                        new_run.font.italic = first_run.font.italic
                                    else:
                                        paragraph.text = text

            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            logger.info(f"Document generated successfully from template: {template_name}")
            return buffer.getvalue()

        except Exception as e:
            logger.error(f"Error generating document from template {template_name}: {str(e)}")
            raise


# Singleton instance
_document_template_service = None

def get_document_template_service() -> DocumentTemplateService:
    """Get or create the document template service singleton."""
    global _document_template_service
    if _document_template_service is None:
        _document_template_service = DocumentTemplateService()
    return _document_template_service
